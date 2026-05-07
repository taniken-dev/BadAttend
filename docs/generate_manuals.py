"""
BadAttend 役職別マニュアル生成スクリプト
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# ヘルパー関数
# ============================================================

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_info_box(doc, text, bg_hex='E8F4FD', border_color='2196F3'):
    """情報ボックス（枠付き段落）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # 左インデント
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    # 枠線を pPr に追加
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '12')
        bdr.set(qn('w:space'), '4')
        bdr.set(qn('w:color'), border_color)
        pBdr.append(bdr)
    pPr.append(pBdr)
    return p

def add_step_table(doc, steps):
    """手順テーブル（番号 + 説明）"""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # ヘッダー
    hdr = table.rows[0].cells
    hdr[0].text = '手順'
    hdr[1].text = '操作内容'
    set_cell_bg(hdr[0], '1565C0')
    set_cell_bg(hdr[1], '1565C0')
    for cell in hdr:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)
    # 列幅
    table.columns[0].width = Cm(2.5)
    table.columns[1].width = Cm(13)

    for i, (step_title, step_desc) in enumerate(steps, 1):
        row = table.add_row().cells
        row[0].text = f'Step {i}\n{step_title}'
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for para in row[0].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        set_cell_bg(row[0], 'E3F2FD')
        row[1].text = step_desc
        for para in row[1].paragraphs:
            para.paragraph_format.left_indent = Cm(0.2)
            for run in para.runs:
                run.font.size = Pt(10.5)
    return table

def add_two_col_table(doc, headers, rows, header_color='1565C0'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], header_color)
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10.5)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            for para in row_cells[i].paragraphs:
                para.paragraph_format.left_indent = Cm(0.2)
                for run in para.runs:
                    run.font.size = Pt(10.5)
    return table

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_page_break(doc):
    doc.add_page_break()

def set_doc_defaults(doc):
    style = doc.styles['Normal']
    style.font.name = 'メイリオ'
    style.font.size = Pt(10.5)
    # A4 サイズ・余白設定
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

def add_cover(doc, title, subtitle, role_label, role_color):
    """表紙ページ"""
    # スペース
    for _ in range(6):
        doc.add_paragraph()
    # アプリ名
    app_p = doc.add_paragraph()
    app_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    app_run = app_p.add_run('BadAttend')
    app_run.font.size = Pt(28)
    app_run.font.bold = True
    app_run.font.color.rgb = RGBColor(*bytes.fromhex('1565C0'))

    # サブタイトル
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run('バドミントン部 出欠管理システム')
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(*bytes.fromhex('455A64'))

    doc.add_paragraph()

    # 役職バッジ
    badge_p = doc.add_paragraph()
    badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge_run = badge_p.add_run(f'【 {role_label} 向けマニュアル 】')
    badge_run.font.size = Pt(18)
    badge_run.font.bold = True
    badge_run.font.color.rgb = RGBColor(*bytes.fromhex(role_color))

    doc.add_paragraph()

    # タイトル
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    title_run.font.size = Pt(16)
    title_run.font.bold = True

    doc.add_paragraph()

    # 日付
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run('2026年5月')
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(*bytes.fromhex('78909C'))

    add_page_break(doc)


# ============================================================
# 1. 部員（member）マニュアル
# ============================================================
def create_member_manual():
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc,
              'はじめてでもわかる！出欠連絡マニュアル',
              'バドミントン部 出欠管理システム',
              '部員', '1565C0')

    # ---- 目次 ----
    add_heading(doc, '目次', level=1)
    toc_items = [
        '1. このアプリでできること',
        '2. はじめてログインする',
        '3. ホーム画面の見かた',
        '4. 出欠連絡のしかた（最も重要！）',
        '5. 出欠ステータスとポイントの仕組み',
        '6. ランキングの見かた',
        '7. 匿名でご意見を送る（意見箱）',
        '8. よくある質問（FAQ）',
    ]
    for item in toc_items:
        add_bullet(doc, item)

    add_page_break(doc)

    # ---- 1. できること ----
    add_heading(doc, '1. このアプリでできること', level=1, color='1565C0')
    add_body(doc, 'BadAttend（バドアテンド）は、千葉工業大学バドミントン部の出欠連絡・成績確認をスマホやパソコンから行えるシステムです。')
    add_two_col_table(doc,
        ['機能', '説明'],
        [
            ['出欠連絡', '練習への出欠・遅刻・欠席を事前に連絡できます'],
            ['ホーム', '今日の練習状況・自分の出席率・直近の出欠をまとめて確認できます'],
            ['カレンダー', '過去・今後の練習予定と出欠状況を月単位で確認できます'],
            ['ランキング', '出席率ランキングでチームの活動状況を確認できます'],
            ['意見箱', '部への意見・要望を匿名で送ることができます'],
        ]
    )
    doc.add_paragraph()

    # ---- 2. ログイン ----
    add_heading(doc, '2. はじめてログインする', level=1, color='1565C0')
    add_info_box(doc, 'このアプリへのログインは LINE アカウントを使用します。LINEアプリがインストールされたスマートフォンをご用意ください。', 'FFF9C4', 'F9A825')
    doc.add_paragraph()

    add_heading(doc, '2-1. はじめてアカウントを作る', level=2)
    add_step_table(doc, [
        ('サイトを開く', '管理者から教えてもらったURLをブラウザで開きます。'),
        ('LINEでログインをタップ', 'ログイン画面の「LINEでログイン」ボタンをタップします。'),
        ('LINE認証', 'LINEアプリが開き、ログイン許可を求められます。「許可する」をタップします。'),
        ('情報入力', '初回のみ、氏名・学年を入力してプロフィールを作成します。'),
        ('承認待ち', '「承認待ち」画面が表示されます。管理者が承認するまでしばらくお待ちください。'),
        ('利用開始', '承認後、自動的にホーム画面に移動してアプリが使えるようになります。'),
    ])
    doc.add_paragraph()

    add_heading(doc, '2-2. 2回目以降のログイン', level=2)
    add_step_table(doc, [
        ('サイトを開く', 'アプリのURLをブラウザで開きます（ブックマーク登録を推奨）。'),
        ('LINEでログインをタップ', '「LINEでログイン」ボタンをタップします。'),
        ('自動ログイン', 'LINEで認証済みの場合は自動的にログインされてホーム画面が開きます。'),
    ])
    doc.add_paragraph()

    add_page_break(doc)

    # ---- 3. ホーム ----
    add_heading(doc, '3. ホーム画面の見かた', level=1, color='1565C0')
    add_body(doc, 'ログイン後に最初に表示されるのがホーム画面です。以下の情報が確認できます。')
    add_two_col_table(doc,
        ['表示エリア', '内容'],
        [
            ['今日の練習', '今日の練習の参加予定者・欠席者の一覧が表示されます'],
            ['自分の出席率', '確定済みの練習に対する自分の出席率（%）が表示されます'],
            ['直近の出欠', '直近10回分の出欠を色付きドットで確認できます（緑＝出席、黄＝遅刻、赤＝欠席）'],
            ['ランキング', '出席率上位3名の表彰台と全員のランキングが表示されます'],
        ]
    )
    doc.add_paragraph()
    add_info_box(doc, '📱 スマホではナビゲーションが画面下部に表示されます。パソコンでは画面上部に表示されます。', 'E8F4FD', '2196F3')

    add_page_break(doc)

    # ---- 4. 出欠連絡 ----
    add_heading(doc, '4. 出欠連絡のしかた（最も重要！）', level=1, color='D32F2F')
    add_info_box(doc, '⚠️ 練習への出欠は必ず事前に連絡しましょう！無連絡欠席はポイントが大きく減点されます。', 'FFEBEE', 'D32F2F')
    doc.add_paragraph()

    add_heading(doc, '4-1. 連絡できる時間帯', level=2)
    add_two_col_table(doc,
        ['曜日', '連絡できる内容'],
        [
            ['月曜日・火曜日（火曜23:59まで）', '今週の水・木・金曜日の練習を事前登録できます'],
            ['水曜日・木曜日・金曜日', '当日の練習のみ登録できます'],
            ['土曜日・日曜日', '出欠連絡はできません（翌週の月曜から登録可能）'],
        ]
    )
    doc.add_paragraph()
    add_info_box(doc, '🔔 練習開始の1時間以内に欠席連絡をすると「緊急欠席」扱いとなり、ポイントが通常より多く減点されます。なるべく早めに連絡しましょう！', 'FFF3E0', 'F57C00')
    doc.add_paragraph()

    add_heading(doc, '4-2. 出欠を登録する手順', level=2)
    add_step_table(doc, [
        ('出欠連絡を開く', 'ナビゲーションの「出欠連絡」をタップします。'),
        ('日付を確認', '登録したい練習日が表示されています。日付ごとにカードが並んでいます。'),
        ('ステータスを選ぶ', '「出席」「遅刻」「欠席」から選択します。'),
        ('欠席理由を選ぶ', '欠席・遅刻の場合は理由（授業・体調不良・私用など）を選択します。'),
        ('登録する', '「登録する」ボタンをタップして完了です。'),
    ])
    doc.add_paragraph()

    add_heading(doc, '4-3. 登録内容を変更する', level=2)
    add_body(doc, '登録後でも、練習開始前であれば内容を変更できます。出欠連絡画面から同じ練習日のカードをタップして、新しい内容を選び直してください。')
    doc.add_paragraph()

    add_heading(doc, '4-4. 体調不良で欠席したとき', level=2)
    add_info_box(doc, '🤒 欠席理由に「体調不良」を選んだ場合、翌日は自動的に「休養推奨モード」になり、出欠連絡ができなくなります（出席できません）。無理せず休みましょう。', 'E8F5E9', '388E3C')

    add_page_break(doc)

    # ---- 5. ポイント ----
    add_heading(doc, '5. 出欠ステータスとポイントの仕組み', level=1, color='1565C0')
    add_body(doc, 'ポイントは初期値1000点からスタートします。欠席の種類によって減点されます。')
    add_two_col_table(doc,
        ['ステータス', '説明', 'ポイント変動'],
        [
            ['出席 ✅', '練習に出席した', '±0'],
            ['遅刻 🕐', '練習に遅刻した', '−10点'],
            ['欠席 📋', '事前に欠席連絡をした', '−20点'],
            ['緊急欠席 ⚡', '練習開始1時間以内に欠席連絡した', '−50点'],
            ['無連絡欠席 ❌', '連絡なしに欠席した（自動記録）', '−100点'],
        ]
    )
    doc.add_paragraph()
    add_info_box(doc, '💡 ポイントはランキングには直接影響しません。ランキングは「出席率」と「技術ランク」で決まります。ポイントは部内での評価指標の一つです。', 'E8F4FD', '2196F3')

    add_page_break(doc)

    # ---- 6. ランキング ----
    add_heading(doc, '6. ランキングの見かた', level=1, color='1565C0')
    add_body(doc, 'ホーム画面でランキングを確認できます。ランキングは「実績確定済み」の練習データをもとに計算されます。')
    add_two_col_table(doc,
        ['項目', '説明'],
        [
            ['出席率', '確定済み練習のうち、出席・遅刻した割合（遅刻は0.5回分としてカウント）'],
            ['選考スコア', '（技術ランク ÷ 6）× 出席率 × 100 で計算されるスコア'],
            ['選考ランク', 'S / A / B / C / D / E で出席率を段階評価'],
        ]
    )
    doc.add_paragraph()
    add_info_box(doc, '🏆 表彰台には出席率トップ3が表示されます。上を目指してコンスタントに参加しましょう！', 'FFF9C4', 'F9A825')

    add_page_break(doc)

    # ---- 7. 意見箱 ----
    add_heading(doc, '7. 匿名でご意見を送る（意見箱）', level=1, color='1565C0')
    add_body(doc, '部への意見・要望・不満などを匿名で送ることができます。誰が送ったかはシステムに記録されません。')
    add_step_table(doc, [
        ('意見箱を開く', 'ナビゲーションの「ご意見箱」アイコンをタップします（封筒マーク）。'),
        ('タイトルを入力', '件名（例：「練習内容について」）を入力します。'),
        ('内容を入力', '具体的な意見・要望を入力します。'),
        ('送信', '「送信する」ボタンをタップして完了です。'),
    ])
    doc.add_paragraph()
    add_info_box(doc, '🔒 送信者情報は一切記録されません。安心して意見を送ってください。', 'E8F5E9', '388E3C')

    add_page_break(doc)

    # ---- 8. FAQ ----
    add_heading(doc, '8. よくある質問（FAQ）', level=1, color='1565C0')

    faqs = [
        ('Q. 出欠の登録を間違えた！', '練習開始前であれば変更できます。出欠連絡画面から同じ日付のカードを選び直してください。練習開始後の修正はマネージャーまたは管理者にお知らせください。'),
        ('Q. 「承認待ち」のまま進めない...', 'アカウント登録後、管理者の承認が必要です。少し時間をおいて再ログインしてみてください。解決しない場合は管理者に連絡してください。'),
        ('Q. ランキングが更新されない', 'ランキングはマネージャーが「実績確定」を行った後に更新されます。練習後しばらく待ってみてください。'),
        ('Q. LINEログインでエラーになる', 'LINEアプリが最新バージョンか確認してください。解決しない場合は管理者に連絡してください。'),
        ('Q. ダークモードにしたい', 'ナビゲーションの「その他」→「ダークモード」から切り替えられます。'),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run_q = p.add_run(q)
        run_q.font.bold = True
        run_q.font.size = Pt(11)
        run_q.font.color.rgb = RGBColor(*bytes.fromhex('1565C0'))
        p2 = doc.add_paragraph(a)
        for run in p2.runs:
            run.font.size = Pt(10.5)
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(8)

    path = os.path.join(OUTPUT_DIR, '部員向けマニュアル.docx')
    doc.save(path)
    print(f'OK 部員向けマニュアル.docx を保存しました')
    return path


# ============================================================
# 2. マネージャー（manager）マニュアル
# ============================================================
def create_manager_manual():
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc,
              '出欠実績管理マニュアル',
              'バドミントン部 出欠管理システム',
              'マネージャー', '2E7D32')

    add_heading(doc, '目次', level=1)
    toc_items = [
        '1. マネージャーの役割',
        '2. ログインとダッシュボードの確認',
        '3. 出欠実績の確定（最重要作業）',
        '4. 出欠データの修正',
        '5. カレンダーでの管理',
        '6. ランキングと選考スコアの確認',
        '7. 注意事項',
    ]
    for item in toc_items:
        add_bullet(doc, item)
    add_page_break(doc)

    # ---- 1. 役割 ----
    add_heading(doc, '1. マネージャーの役割', level=1, color='2E7D32')
    add_body(doc, 'マネージャーは、部員の出欠を正確な「実績」として確定させる重要な役割を担います。')
    add_two_col_table(doc,
        ['できること', '説明'],
        [
            ['出欠実績の確定', '練習後に全部員の出欠を確定させます。これによりランキングが更新されます'],
            ['出欠データの修正', '部員が誤って登録した出欠を修正できます'],
            ['カレンダー管理', '練習セッションの詳細確認・出欠状況の一覧確認ができます'],
            ['全部員の出欠閲覧', '全部員の出欠記録を確認できます'],
        ],
        header_color='2E7D32'
    )
    doc.add_paragraph()
    add_info_box(doc, '⚠️ マネージャーはメンバーの承認・削除・ロール変更はできません。これらは管理者（admin）の権限です。', 'FFF3E0', 'F57C00')

    add_page_break(doc)

    # ---- 2. ダッシュボード ----
    add_heading(doc, '2. ログインとダッシュボードの確認', level=1, color='2E7D32')
    add_body(doc, 'LINEアカウントでログインします。管理者にマネージャー権限を付与してもらうと、追加の操作メニューが表示されます。')
    add_two_col_table(doc,
        ['確認項目', '場所'],
        [
            ['今日の参加予定者・欠席者', 'ホーム画面'],
            ['全体の出席率・ランキング', 'ホーム画面'],
            ['月次の練習カレンダー', 'カレンダー画面'],
            ['実績確定ボタン', 'カレンダー → 日付クリック → セッション詳細'],
        ],
        header_color='2E7D32'
    )

    add_page_break(doc)

    # ---- 3. 実績確定 ----
    add_heading(doc, '3. 出欠実績の確定（最重要作業）', level=1, color='D32F2F')
    add_info_box(doc, '🔑 実績確定を行うまでランキングに反映されません。練習終了後できるだけ早く実施してください。', 'FFEBEE', 'D32F2F')
    doc.add_paragraph()

    add_heading(doc, '3-1. 実績確定の手順', level=2)
    add_step_table(doc, [
        ('カレンダーを開く', 'ナビゲーションの「カレンダー」をタップします。'),
        ('練習日をタップ', '実績確定したい練習日（水・木・金）をタップします。'),
        ('出欠状況を確認', 'セッション詳細画面で全部員の出欠状況を確認します。必要であればこの時点で修正を行います。'),
        ('実績確定ボタン', '「実績を確定する」ボタンをタップします。'),
        ('確認ダイアログ', '「本当に確定しますか？」というダイアログが表示されます。内容を確認して「確定する」をタップします。'),
        ('完了', '確定が完了すると、ランキングが自動的に更新されます。確定者名と日時が記録されます。'),
    ])
    doc.add_paragraph()

    add_heading(doc, '3-2. 実績確定後の修正について', level=2)
    add_info_box(doc, '⚠️ 実績確定後でもデータの修正は可能ですが、確定を取り消すことも一部できます。修正が必要な場合は管理者に相談してください。', 'FFF3E0', 'F57C00')
    doc.add_paragraph()

    add_heading(doc, '3-3. 実績確定の対象ステータス', level=2)
    add_body(doc, '実績確定では、部員が事前に連絡した「予定ステータス」を、実際に起きた「実績ステータス」に変換します。')
    add_two_col_table(doc,
        ['事前連絡（status）', '実績確定後（result_status）', '意味'],
        [
            ['出席', '出席', '予定通り出席'],
            ['遅刻', '遅刻', '予定通り遅刻'],
            ['欠席', '欠席', '事前連絡あり欠席'],
            ['緊急欠席', '緊急欠席', '1時間以内の連絡欠席'],
            ['（連絡なし）', '無連絡欠席', '当日連絡なし → 自動的に記録'],
        ],
        header_color='2E7D32'
    )

    add_page_break(doc)

    # ---- 4. 修正 ----
    add_heading(doc, '4. 出欠データの修正', level=1, color='2E7D32')
    add_body(doc, '部員が誤って登録した出欠や、当日の実際の状況と異なる記録を修正できます。')
    add_step_table(doc, [
        ('カレンダーを開く', '修正したい練習日をカレンダーからタップします。'),
        ('部員を選択', 'セッション詳細で修正したい部員の出欠行をタップします。'),
        ('ステータス変更', '正しいステータスに変更します。'),
        ('保存', '変更を保存します。変更内容は自動的にポイントに反映されます。'),
    ])
    doc.add_paragraph()
    add_info_box(doc, '💡 ポイントの変動はデータベースのトリガーで自動計算されます。手動で計算する必要はありません。', 'E8F4FD', '2196F3')

    add_page_break(doc)

    # ---- 5. カレンダー ----
    add_heading(doc, '5. カレンダーでの管理', level=1, color='2E7D32')
    add_two_col_table(doc,
        ['操作', '説明'],
        [
            ['月移動', '画面上部の矢印で前月・翌月に移動できます'],
            ['日付タップ', '練習がある日付をタップするとセッション詳細が表示されます'],
            ['実績確定済み表示', '確定済みの練習日にはチェックマークが表示されます'],
            ['練習中止確認', '中止になった練習はカレンダー上で異なる表示になります'],
        ],
        header_color='2E7D32'
    )

    add_page_break(doc)

    # ---- 6. ランキング ----
    add_heading(doc, '6. ランキングと選考スコアの確認', level=1, color='2E7D32')
    add_body(doc, 'ホーム画面のランキングは実績確定後に自動更新されます。')
    add_two_col_table(doc,
        ['指標', '計算式', '備考'],
        [
            ['出席率', '(出席数 + 遅刻数×0.5) ÷ 確定済みセッション数 × 100', '%表示'],
            ['選考スコア', '(技術ランク ÷ 6) × 出席率 × 100', '技術ランクは管理者が設定'],
            ['選考ランク', 'S/A/B/C/D/E（出席率で自動判定）', '自動計算'],
        ],
        header_color='2E7D32'
    )
    doc.add_paragraph()

    # ---- 7. 注意事項 ----
    add_heading(doc, '7. 注意事項', level=1, color='D32F2F')
    notes = [
        '実績確定は練習終了後、できるだけ当日中に行ってください。',
        '確定前に出欠状況をよく確認してください。確定後の修正は手間がかかります。',
        '無連絡欠席の部員がいる場合、実績確定時に自動的に「無連絡欠席」ステータスが付きます。',
        '体調不良で欠席した部員は翌日自動的に「休養推奨モード」になり出欠連絡ができません。',
        'マネージャー権限では部員の削除・承認はできません。管理者に依頼してください。',
    ]
    for note in notes:
        add_bullet(doc, f'⚠️ {note}')

    path = os.path.join(OUTPUT_DIR, 'マネージャー向けマニュアル.docx')
    doc.save(path)
    print(f'OK マネージャー向けマニュアル.docx を保存しました')
    return path


# ============================================================
# 3. 管理者（admin）マニュアル
# ============================================================
def create_admin_manual():
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc,
              'システム管理者マニュアル',
              'バドミントン部 出欠管理システム',
              '管理者', 'B71C1C')

    add_heading(doc, '目次', level=1)
    toc_items = [
        '1. 管理者の役割と権限',
        '2. メンバー管理',
        '3. 新規部員の承認手順',
        '4. 部員情報の編集',
        '5. 部員の削除',
        '6. 技術ランクの設定',
        '7. 警告・ペナルティ管理',
        '8. 意見箱の確認',
        '9. セキュリティと権限について',
        '10. トラブルシューティング',
    ]
    for item in toc_items:
        add_bullet(doc, item)
    add_page_break(doc)

    # ---- 1. 役割 ----
    add_heading(doc, '1. 管理者の役割と権限', level=1, color='B71C1C')
    add_body(doc, '管理者（admin）はシステム全体を管理する最高権限を持ちます。マネージャーの全機能に加えて、以下の機能が利用できます。')
    add_two_col_table(doc,
        ['権限', '説明'],
        [
            ['メンバー承認', '新規登録した部員を承認してアクセスを許可します'],
            ['ロール変更', '部員の役職（member/manager/admin/coach）を変更します'],
            ['技術ランク設定', '各部員の技術ランク（E〜S級）を設定します'],
            ['部員削除', 'アカウントを完全に削除します（取り消し不可）'],
            ['意見箱確認', '匿名で送られた部員からの意見を確認します'],
            ['警告フラグ管理', '部費未払い・品位違反などの警告を記録します'],
            ['マネージャー全機能', '実績確定・出欠修正も含む全機能が利用可能'],
        ],
        header_color='B71C1C'
    )
    doc.add_paragraph()
    add_info_box(doc, '🔐 管理者権限は最小限の人数に留めてください。権限の乱用はシステムの信頼性を損ないます。', 'FFEBEE', 'D32F2F')

    add_page_break(doc)

    # ---- 2. メンバー管理画面 ----
    add_heading(doc, '2. メンバー管理', level=1, color='B71C1C')
    add_body(doc, 'ナビゲーションの「メンバー管理」から開きます。画面は以下のセクションで構成されています。')
    add_two_col_table(doc,
        ['セクション', '内容'],
        [
            ['承認待ち一覧', '新規登録してまだ承認されていない部員の一覧'],
            ['全メンバー一覧', '承認済みの全部員。学年・技術ランク・役職の編集が可能'],
            ['孤立ユーザー', 'プロフィールが未作成のアカウント（異常状態の検出）'],
        ],
        header_color='B71C1C'
    )

    add_page_break(doc)

    # ---- 3. 承認 ----
    add_heading(doc, '3. 新規部員の承認手順', level=1, color='B71C1C')
    add_info_box(doc, '📋 新入部員が登録すると「承認待ち」に表示されます。承認するまでアプリを使えません。速やかに対応してください。', 'E8F4FD', '2196F3')
    doc.add_paragraph()
    add_step_table(doc, [
        ('メンバー管理を開く', 'ナビゲーションの「メンバー管理」をタップします。'),
        ('承認待ち確認', '画面上部の「承認待ち」セクションに新規登録者が表示されています。'),
        ('情報確認', '氏名・学年を確認します。本人確認が取れている場合のみ承認してください。'),
        ('承認ボタン', '「承認する」ボタンをタップします。'),
        ('完了', '承認完了後、部員はすぐにアプリを使えるようになります。'),
    ])
    doc.add_paragraph()
    add_info_box(doc, '⚠️ 不審なアカウントは承認しないでください。削除する場合は「削除」ボタンから対応できます。', 'FFF3E0', 'F57C00')

    add_page_break(doc)

    # ---- 4. 部員情報編集 ----
    add_heading(doc, '4. 部員情報の編集', level=1, color='B71C1C')
    add_body(doc, '全メンバー一覧から各部員の編集アイコンをタップすると、以下の情報を変更できます。')
    add_two_col_table(doc,
        ['編集項目', '説明', '注意点'],
        [
            ['表示名', 'アプリ上に表示される名前', 'スペースは「_」区切りなど統一推奨'],
            ['学年', '1〜4年（または卒業）', '毎年4月に更新してください'],
            ['技術ランク', 'E/D/C/B/A/S の6段階', '選考スコアの計算に使用されます'],
            ['役職（ロール）', 'member/manager/admin/coach', '変更は慎重に。admin は最小限に'],
        ],
        header_color='B71C1C'
    )
    doc.add_paragraph()

    add_heading(doc, '役職（ロール）の説明', level=2)
    add_two_col_table(doc,
        ['役職', '主な権限'],
        [
            ['member（部員）', '自分の出欠連絡・ランキング閲覧・意見箱投稿'],
            ['manager（マネージャー）', '+全部員の出欠実績確定・出欠修正'],
            ['admin（管理者）', '+メンバー承認・削除・ロール変更・意見確認・警告管理'],
            ['coach（顧問）', 'メンバー一覧閲覧のみ（出欠連絡不可）'],
        ],
        header_color='B71C1C'
    )

    add_page_break(doc)

    # ---- 5. 削除 ----
    add_heading(doc, '5. 部員の削除', level=1, color='B71C1C')
    add_info_box(doc, '🚨 削除は取り消せません。出欠履歴・ポイント履歴を含むすべてのデータが削除されます。必ず本人に確認を取ってから行ってください。', 'FFEBEE', 'D32F2F')
    doc.add_paragraph()
    add_step_table(doc, [
        ('メンバー管理を開く', 'ナビゲーションの「メンバー管理」をタップします。'),
        ('対象部員を探す', '全メンバー一覧から削除したい部員を探します。'),
        ('削除ボタン', '該当部員の「削除」ボタンをタップします。'),
        ('確認ダイアログ', '「本当に削除しますか？この操作は取り消せません」というダイアログが表示されます。'),
        ('実行', '確認の上、「削除する」をタップします。'),
    ])

    add_page_break(doc)

    # ---- 6. 技術ランク ----
    add_heading(doc, '6. 技術ランクの設定', level=1, color='B71C1C')
    add_body(doc, '技術ランクは選考スコアの計算に使用される重要な指標です。定期的に更新してください。')
    add_two_col_table(doc,
        ['ランク', '目安'],
        [
            ['S級（6）', '全国・関東大会レベル'],
            ['A級（5）', '県大会上位レベル'],
            ['B級（4）', '県大会出場レベル'],
            ['C級（3）', '地区大会レベル'],
            ['D級（2）', '部内中級レベル'],
            ['E級（1）', '初心者・入門レベル'],
        ],
        header_color='B71C1C'
    )
    doc.add_paragraph()
    add_info_box(doc, '💡 選考スコア = (技術ランク ÷ 6) × 出席率 × 100。ランクが高いほどスコアが高くなります。部員との合意の上で設定してください。', 'E8F4FD', '2196F3')

    add_page_break(doc)

    # ---- 7. 警告 ----
    add_heading(doc, '7. 警告・ペナルティ管理', level=1, color='B71C1C')
    add_body(doc, 'ルール違反や部費未払いなどの問題行動に対して警告フラグを記録できます。')
    add_two_col_table(doc,
        ['警告タイプ', '説明'],
        [
            ['dues_overdue（部費未払い）', '部費の支払いが遅延している'],
            ['absent_no_report（無連絡欠席継続）', '無連絡欠席が続いている'],
            ['conduct（品位違反）', '部の品位を損なう行為'],
        ],
        header_color='B71C1C'
    )
    doc.add_paragraph()
    add_two_col_table(doc,
        ['重大度（severity）', '説明'],
        [
            ['warning（警告）', '最初の警告'],
            ['final_warning（最終警告）', '改善が見られない場合'],
            ['expelled（除籍）', '最終的な措置'],
        ],
        header_color='B71C1C'
    )
    doc.add_paragraph()
    add_info_box(doc, '📢 警告中の部員はランキング画面に ⚠️ アイコンが表示されます（管理者・マネージャーのみ表示）。', 'E8F4FD', '2196F3')

    add_page_break(doc)

    # ---- 8. 意見箱 ----
    add_heading(doc, '8. 意見箱の確認', level=1, color='B71C1C')
    add_body(doc, '部員から匿名で送られた意見を確認できます。管理者のみがアクセスできる画面です。')
    add_step_table(doc, [
        ('意見確認を開く', 'ナビゲーションの「その他（ギアアイコン）」→「届いた意見を確認」をタップします。'),
        ('意見一覧', '届いた意見がタイトルと投稿日時と共に一覧表示されます。'),
        ('内容確認', '各意見をタップして本文を読みます。'),
        ('対応', '必要に応じてミーティングや掲示板などで部員に共有・回答します。'),
    ])
    doc.add_paragraph()
    add_info_box(doc, '🔒 誰が送ったかはシステムに記録されていません。匿名性を守るため、内容から個人を特定しようとしないでください。', 'E8F5E9', '388E3C')

    add_page_break(doc)

    # ---- 9. セキュリティ ----
    add_heading(doc, '9. セキュリティと権限について', level=1, color='B71C1C')
    add_body(doc, 'このシステムはRow Level Security（RLS）によってデータが保護されています。')
    add_two_col_table(doc,
        ['データ', 'アクセス権限'],
        [
            ['プロフィール情報', '承認済み部員全員が閲覧可能、本人のみ更新可能'],
            ['出欠記録', '本人のみ読み書き可能、manager/adminは全員分更新可能'],
            ['ポイント履歴', '本人のみ閲覧可能'],
            ['警告フラグ', 'admin/managerのみ読み書き可能'],
            ['意見箱', '全員が投稿可能、adminのみ閲覧可能'],
        ],
        header_color='B71C1C'
    )
    doc.add_paragraph()

    # ---- 10. トラブルシューティング ----
    add_heading(doc, '10. トラブルシューティング', level=1, color='B71C1C')
    faqs = [
        ('Q. 部員が「承認待ち」から進めないと言っている',
         'メンバー管理画面の「承認待ち」セクションを確認し、承認ボタンをタップしてください。'),
        ('Q. ランキングが更新されない',
         'マネージャーが実績確定を行っていない可能性があります。カレンダーから未確定の練習日を確認してください。'),
        ('Q. 部員がログインできない',
         '①承認済みか確認 ②LINEアプリが最新バージョンか部員に確認してもらう ③それでも解決しない場合はSupabase管理画面でOAuth設定を確認'),
        ('Q. プロフィールが作成されていない部員がいる',
         'メンバー管理の「孤立ユーザー」セクションに表示されます。手動でプロフィール作成が必要です。'),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run_q = p.add_run(q)
        run_q.font.bold = True
        run_q.font.size = Pt(11)
        run_q.font.color.rgb = RGBColor(*bytes.fromhex('B71C1C'))
        p2 = doc.add_paragraph(a)
        for run in p2.runs:
            run.font.size = Pt(10.5)
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(8)

    path = os.path.join(OUTPUT_DIR, '管理者向けマニュアル.docx')
    doc.save(path)
    print(f'OK 管理者向けマニュアル.docx を保存しました')
    return path


# ============================================================
# 4. 顧問（coach）マニュアル
# ============================================================
def create_coach_manual():
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc,
              '顧問向け閲覧マニュアル',
              'バドミントン部 出欠管理システム',
              '顧問', '4A148C')

    add_heading(doc, '目次', level=1)
    toc_items = [
        '1. 顧問アカウントについて',
        '2. ログイン方法',
        '3. 閲覧できる情報',
        '4. 注意事項',
    ]
    for item in toc_items:
        add_bullet(doc, item)
    add_page_break(doc)

    # ---- 1. 顧問について ----
    add_heading(doc, '1. 顧問アカウントについて', level=1, color='4A148C')
    add_body(doc, '顧問（coach）アカウントは、部の活動状況を把握するための「閲覧専用」アカウントです。')
    add_info_box(doc, '📌 顧問アカウントでは出欠連絡・データの変更・部員管理などの操作は行えません。閲覧のみに限定されています。', 'EDE7F6', '7B1FA2')
    doc.add_paragraph()
    add_two_col_table(doc,
        ['できること', 'できないこと'],
        [
            ['部員一覧の閲覧', '出欠連絡'],
            ['出席率・ランキングの確認', '出欠データの変更'],
            ['練習カレンダーの確認', '部員の承認・削除'],
            ['今日の参加状況の確認', '技術ランクの変更'],
            ['', '意見箱の閲覧'],
        ],
        header_color='4A148C'
    )

    add_page_break(doc)

    # ---- 2. ログイン ----
    add_heading(doc, '2. ログイン方法', level=1, color='4A148C')
    add_info_box(doc, 'このアプリへのログインは LINE アカウントを使用します。LINEアプリがインストールされたスマートフォンをご用意ください。', 'EDE7F6', '7B1FA2')
    doc.add_paragraph()
    add_step_table(doc, [
        ('サイトを開く', '管理者（部長）から教えてもらったURLをブラウザで開きます（ブックマーク登録を推奨）。'),
        ('LINEでログインをタップ', '「LINEでログイン」ボタンをタップします。'),
        ('LINE認証', 'LINEアプリが開き、ログイン許可を求められます。「許可する」をタップします。'),
        ('自動ログイン', '認証が完了すると自動的にアプリに戻り、ホーム画面が表示されます。'),
    ])
    doc.add_paragraph()
    add_info_box(doc, 'ログインできない場合は管理者（部長）にお問い合わせください。', 'EDE7F6', '7B1FA2')

    add_page_break(doc)

    # ---- 3. 閲覧できる情報 ----
    add_heading(doc, '3. 閲覧できる情報', level=1, color='4A148C')

    add_heading(doc, '3-1. ホーム画面', level=2)
    add_body(doc, 'ログイン後のトップ画面で以下を確認できます。')
    add_bullet(doc, '今日の練習に参加予定の部員・欠席予定の部員一覧')
    add_bullet(doc, '各部員の出席率ランキング（表彰台形式）')
    add_bullet(doc, '全部員のランキング一覧（出席率・選考スコア・選考ランク）')
    doc.add_paragraph()

    add_heading(doc, '3-2. カレンダー画面', level=2)
    add_body(doc, '月単位で練習予定と出欠状況を確認できます。')
    add_bullet(doc, '各練習日の出席者・欠席者の人数')
    add_bullet(doc, '実績確定済みかどうかの状態')
    add_bullet(doc, '練習の中止・変更の有無')
    doc.add_paragraph()

    add_heading(doc, '3-3. 部員一覧', level=2)
    add_body(doc, 'メンバー管理画面で部員の基本情報を閲覧できます。')
    add_bullet(doc, '氏名・学年・技術ランク・役職')
    add_bullet(doc, 'LINEアイコン画像（設定している場合）')
    doc.add_paragraph()

    add_heading(doc, '3-4. 選考スコアの見かた', level=2)
    add_two_col_table(doc,
        ['項目', '内容'],
        [
            ['出席率', '確定済み練習に対して出席・遅刻した割合（%）'],
            ['選考スコア', '技術ランクと出席率を組み合わせたスコア（最高100点）'],
            ['選考ランク', 'S/A/B/C/D/E の6段階評価'],
        ],
        header_color='4A148C'
    )

    add_page_break(doc)

    # ---- 4. 注意事項 ----
    add_heading(doc, '4. 注意事項', level=1, color='4A148C')
    notes = [
        '閲覧した部員情報は部外秘です。外部に漏洩しないよう管理してください。',
        'ランキングはあくまで参考指標です。実際の選考は顧問・部員全体で話し合って決定してください。',
        'ログインしたままパソコン・スマホを放置しないようにしてください。使用後は必ずログアウトしてください。',
        '操作方法でわからないことがあれば、管理者（部長）にお問い合わせください。',
    ]
    for note in notes:
        add_bullet(doc, f'✅ {note}')

    path = os.path.join(OUTPUT_DIR, '顧問向けマニュアル.docx')
    doc.save(path)
    print(f'OK 顧問向けマニュアル.docx を保存しました')
    return path


# ============================================================
# メイン
# ============================================================
if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    create_member_manual()
    create_manager_manual()
    create_admin_manual()
    create_coach_manual()
    print('\n🎉 全マニュアルの生成が完了しました！')
    print(f'保存先: {OUTPUT_DIR}')
